# -*- coding: utf-8 -*-
"""Render 陈亚运's resume markdown into a compact 2-page .docx."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CJK_FONT = "PingFang SC"
LATIN_FONT = "Calibri"
ACCENT = RGBColor(0x1F, 0x4E, 0x79)   # deep blue
GRAY = RGBColor(0x55, 0x55, 0x55)
DARK = RGBColor(0x22, 0x22, 0x22)

doc = Document()

# ---- page margins (tight, to fit 2 pages) ----
sec = doc.sections[0]
sec.top_margin = Cm(1.2)
sec.bottom_margin = Cm(1.2)
sec.left_margin = Cm(1.6)
sec.right_margin = Cm(1.6)

# ---- default style ----
normal = doc.styles["Normal"]
normal.font.name = LATIN_FONT
normal.font.size = Pt(9)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.line_spacing = 1.02


def set_run(run, size=9, bold=False, color=DARK, latin=LATIN_FONT):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = latin
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), CJK_FONT)


def add_para(space_after=1.5, space_before=0.0, align=None, line=1.02):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line
    if align:
        p.alignment = align
    return p


def bottom_border(p, size=6, color="1F4E79"):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def section(title):
    p = add_para(space_after=2.5, space_before=6)
    r = p.add_run(title)
    set_run(r, size=11.5, bold=True, color=ACCENT)
    bottom_border(p)


def bullet(runs, space_after=1.5, indent=0.35):
    """runs: list of (text, kwargs)"""
    p = add_para(space_after=space_after)
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    rb = p.add_run("· ")
    set_run(rb, size=9, color=ACCENT, bold=True)
    for text, kw in runs:
        r = p.add_run(text)
        set_run(r, **kw)
    return p


# ================= NAME =================
p = add_para(space_after=1, align=WD_ALIGN_PARAGRAPH.LEFT)
r = p.add_run("陈亚运")
set_run(r, size=20, bold=True, color=DARK)

p = add_para(space_after=3)
r = p.add_run("13121917705  ·  1019880993@qq.com  ·  期望城市：北京")
set_run(r, size=9, color=GRAY)
p = add_para(space_after=2)
r = p.add_run("目标岗位：")
set_run(r, size=9, bold=True, color=DARK)
r = p.add_run("Agent 开发工程师、服务端研发工程师")
set_run(r, size=9, color=DARK)

# ================= 教育背景 =================
section("教育背景")
p = add_para(space_after=1.5)
r = p.add_run("北京邮电大学")
set_run(r, size=9.5, bold=True, color=DARK)
r = p.add_run("  计算机科学与技术 · 硕士")
set_run(r, size=9, color=DARK)
r = p.add_run("                                                       2013.09 - 2016.04")
set_run(r, size=9, color=GRAY)

p = add_para(space_after=1.5)
r = p.add_run("安徽工业大学")
set_run(r, size=9.5, bold=True, color=DARK)
r = p.add_run("  网络工程 · 本科")
set_run(r, size=9, color=DARK)
r = p.add_run("                                                                     2008.09 - 2012.06")
set_run(r, size=9, color=GRAY)

# ================= 工作经历 =================
section("工作经历")


def job(company, role, period, desc_runs=None):
    p = add_para(space_after=1.5, space_before=2)
    r = p.add_run(company)
    set_run(r, size=9.5, bold=True, color=DARK)
    r = p.add_run("   " + role)
    set_run(r, size=9, color=ACCENT, bold=True)
    r = p.add_run("   |   " + period)
    set_run(r, size=9, color=GRAY)


job("滴滴 · 网约车开放平台", "资深研发工程师", "2019.07 - 至今")
bullet([("主导鸿鹄 SaaS 司机 DDD 领域化改造 + AI-native 研发提效 Agent 落地（详见项目一）", dict(size=9))])
bullet([("主导鸿鹄 SaaS 司机多产品线架构升级（详见项目二）；主导融合小猪自营司机主流程升级（详见项目三）", dict(size=9))])
bullet([("作为核心研发参与鸿鹄 SaaS 从 0 到 1 建设，参与开放平台日常迭代", dict(size=9))])

job("滴滴 · 顺风车事业部", "高级研发工程师", "2018.05 - 2019.09")
bullet([("参与顺风车交易管控日常迭代", dict(size=9))])

job("百度外卖", "高级研发工程师", "2016.04 - 2018.04")
bullet([("参与百度外卖配送日常迭代", dict(size=9))])

# ================= 项目经验 =================
section("项目经验")


def project(title, meta):
    p = add_para(space_after=1, space_before=3)
    r = p.add_run(title)
    set_run(r, size=10, bold=True, color=DARK)
    p = add_para(space_after=1.5)
    r = p.add_run(meta)
    set_run(r, size=8.5, color=GRAY)


def sub(label, text):
    p = add_para(space_after=1.5)
    r = p.add_run(label + "  ")
    set_run(r, size=9, bold=True, color=ACCENT)
    r = p.add_run(text)
    set_run(r, size=9, color=DARK)


def stack(text):
    p = add_para(space_after=2)
    r = p.add_run(text)
    set_run(r, size=8.5, color=GRAY)


# ---- 项目一 ----
project("项目一：SaaS 司机主流程领域化改造 + AI-native 研发提效 Agent",
        "2025.11 - 至今  |  主导 SaaS 司机架构治理  |  8 人团队")
sub("背景", "SaaS 司机域沉淀 10+ 核心模块、60W+ 行历史代码；传统事务脚本模式、缺少业务抽象，事故率高、开发与新人培养周期长。")
sub("领域建模", "对主流程做事件风暴推演，沉淀 DDD 限界上下文，划分流程单据与司机运力中心，统一依赖架构，作为 Agent 检索与变更影响分析的事实基线。")
sub("开发框架", "制定符合 DDD 的四层代码框架并迁移改造历史代码；自研在线流量 diff 框架，保证模块迁移稳定性。")
sub("AI-native 提效", "基于 Claude Code + harness engineering 自研 workflow 编排框架，驱动“需求理解→代码生成→评审→上线→沉淀知识库”全链路。")
sub("降本与记忆", "自研 MCP 工具支持按模块/文件粒度代码分析，单次任务 token 消耗较原始 prompt 下降 30%+；自研长效记忆与项目知识库，增强跨会话上下文连续性。")
sub("成果", "完成主流程领域梳理，明确限界上下文/领域架构/开发框架；沉淀 AI-native 内部框架与 Skill，技术改造与开发效率提升 20%。")
stack("Go · Claude Code · Skill · harness engineering · loop engineering · MCP · DDD")

# ---- 项目二 ----
project("项目二：SaaS 司机多产品线架构升级",
        "2024.11 - 2025.10  |  主导多产品线架构升级  |  协调 30+ 团队  |  司机核心团队 10+ 人")
sub("背景", "业务需探索顺风车、出租车场景，而 SaaS 司机架构仅支持快车，需快速支撑顺风车、出租车 MVP 上线。")
sub("多产品线架构", "统一在线流程抽象、领域实体抽象、产品线配置化；按产品/城市/司机类型多维灰度，支持单产品线秒级回滚。")
sub("跨团队协同", "协调 30+ 团队完成存量系统架构盘点与新架构落地；建设司机全链路 trace、告警与资金强一致对账。")
sub("成果", "司机在线架构从单产品线升级为同时支撑快车/顺风车/出租车三条产品线，架构升级 0 事故。")
stack("SaaS 司机 · 顺风车 · 出租车 · 多产品线灰度 · 可观测性")

# ---- 项目三 ----
project("项目三：鸿鹄 SaaS 支持小猪自营司机主流程升级",
        "2023.01 - 2024.10  |  主导 SaaS 司机主流程升级  |  10+ 核心团队  |  10w+ 司机 · 100w+ 订单")
sub("背景", "鸿鹄 SaaS 司机与小猪分属两团队但业务高度重叠，研发资源浪费严重；需升级架构以支持自营司机 60+ 项新功能。")
sub("架构梳理", "从业务流程、存储依赖、稳定性三维度对 60+ 新功能分层归纳与改造优先级排序，确定升级整体框架。")
sub("流程重构", "撮合/费用/安全管控流程分两阶段重构（先建新流程接入自营司机，再切 SaaS 司机流量）；7 维度稳定性治理。")
sub("成果", "完成司机融合，业务迭代效率提升 20%+，0 重大事故。")
stack("Go · 撮合引擎 · 费用流程 · 安全管控 · 稳定性治理 · 在线流量迁移")

# ================= 技术栈 =================
section("技术栈")
bullet([("编程语言：", dict(size=9, bold=True, color=ACCENT)), ("PHP、Go，熟悉基本原理、常用框架与组件", dict(size=9))])
bullet([("Agent 框架与工程化：", dict(size=9, bold=True, color=ACCENT)), ("Claude Code、Prompt、Harness、Loop、MCP、Workflow 编排框架", dict(size=9))])
bullet([("分布式与高并发：", dict(size=9, bold=True, color=ACCENT)), ("限流、熔断、降级、单元化、全链路压测、可观测性建设", dict(size=9))])
bullet([("架构治理：", dict(size=9, bold=True, color=ACCENT)), ("DDD 领域驱动设计、复杂业务系统重构、模块解耦、依赖图谱与变更影响分析", dict(size=9))])
bullet([("项目管理：", dict(size=9, bold=True, color=ACCENT)), ("任务分解、排期推动、风险管理、复盘总结", dict(size=9))])
bullet([("基础设施：", dict(size=9, bold=True, color=ACCENT)), ("RPC、Redis、MySQL、DDMQ、ES、Docker", dict(size=9))])

out = "/Users/chenyayun/work/resume/v1/resume_chenyayun_agent_engineer-v1.docx"
doc.save(out)
print("saved:", out)
